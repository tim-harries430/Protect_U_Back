from __future__ import annotations

from copy import deepcopy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


SOURCE = Path(
    r"C:\Users\LXC\OneDrive\Desktop\pub_论文_终稿_实验加术语_摘要Soak_含实验哈希.docx"
)
OUTPUT = Path(
    r"D:\dev\sp\.docx_work\pub_论文_baseline图_引用角标.docx"
)
CHART = Path(r"D:\dev\sp\.docx_work\baseline_security_utility.png")

TITLE = "十一、Baseline 对照：门看见了变化，却还没有学会解释授权"
BODY_1 = (
    "为了知道连续性机制补上了什么，笔者保持 RT01–RT10 的 Win32 注入、v3 oracle、"
    "attack/control 配对与生产模块 SHA 不变，只在同一窗口旁挂五种 baseline 和两种 ablation。"
    "Naked Agent 与 Prompt Guardrail 代表没有物理观察的入口；Keyword Guardrail 代表文本与元数据扫描；"
    "Content SHA 与 Legacy Endpoint 代表单次 enter/exit 端点比较。Access-only 撤去 process witness，"
    "Process-only 撤去 access witness，Full PUB 作为参照。八种方法共同接受 10 条 attack 与 10 条 control。"
)
BODY_2 = (
    "图11把纵轴设为 attack capture，横轴设为 benign completion。Naked Agent 与 Prompt Guardrail 位于右下角："
    "日常动作全过，十条攻击也全过。Content SHA 捕获 8/10，合法完成 2/10；Legacy Endpoint 与 Access-only "
    "捕获 9/10，合法完成 1/10，二者漏掉窗口内出生又消失的 RT09。Keyword Guardrail、Process-only 与 "
    "Full PUB 都来到左上角：攻击捕获 1.000，合法完成 0。八个点沿左上至右下排列，没有一个进入右上方的理想区域。"
)
BODY_3 = (
    "这张图给出两个结论。T 轴的活体拍点补上了端点法的瞬态盲区，Process-only 与 Full PUB 把攻击捕获推至 "
    "1.000；P_auth 仍未消去授权 control，Full PUB 的二分类 F1 为 0.667，与全拦截的 Keyword Guardrail 持平。"
    "因此，本轮 baseline 支持“过程观察增加了可见性”，尚不支持“系统已经取得安全—可用性优势”。后续修订只动 "
    "P_auth 与责任投影，题目、oracle 和 baseline 哈希继续冻结。"
)
CAPTION = (
    "图11　同一 20 条配对注入下的安全—可用性平面。纵轴为攻击捕获率，横轴为合法完成率；"
    "圆点表示 baseline，三角表示消融，星形表示 Full PUB，右上角为理想区域。"
    "数据：baseline_score.json，SHA-256 "
    "d86ba3e3ddc2c82001bdd8ebe5f24a9004a504a37ab1ea927167c0da15f84e98。"
)


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


FONT = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"


def draw_centered(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, fnt, fill) -> None:
    box = draw.textbbox((0, 0), text, font=fnt)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=fnt, fill=fill)


def dashed_line(draw: ImageDraw.ImageDraw, start, end, fill, width=3, dash=14) -> None:
    x1, y1 = start
    x2, y2 = end
    length = ((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5
    if not length:
        return
    ux, uy = (x2 - x1) / length, (y2 - y1) / length
    distance = 0.0
    while distance < length:
        stop = min(distance + dash, length)
        draw.line(
            (x1 + ux * distance, y1 + uy * distance, x1 + ux * stop, y1 + uy * stop),
            fill=fill,
            width=width,
        )
        distance += dash * 1.8


def triangle(draw, center, radius, fill, outline, width=4) -> None:
    x, y = center
    points = ((x, y - radius), (x - radius, y + radius), (x + radius, y + radius))
    draw.polygon(points, fill=fill)
    draw.line((*points, points[0]), fill=outline, width=width, joint="curve")


def star(draw, center, outer, inner, fill, outline, width=4) -> None:
    import math

    x, y = center
    points = []
    for i in range(10):
        angle = -math.pi / 2 + i * math.pi / 5
        radius = outer if i % 2 == 0 else inner
        points.append((x + radius * math.cos(angle), y + radius * math.sin(angle)))
    draw.polygon(points, fill=fill)
    draw.line((*points, points[0]), fill=outline, width=width, joint="curve")


def make_chart(path: Path) -> None:
    width, height = 2200, 1320
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    black = "#1e252b"
    grey = "#66717a"
    light_grey = "#d9dee2"
    blue = "#2867a5"
    red = "#b52b31"
    green = "#dff1e3"
    title_font = font(FONT_BOLD, 54)
    panel_font = font(FONT_BOLD, 34)
    label_font = font(FONT, 27)
    small_font = font(FONT, 24)
    tiny_font = font(FONT, 21)

    draw_centered(draw, (width // 2, 62), "Baseline 与 PUB：安全—可用性对照", title_font, black)

    left, top, right, bottom = 150, 180, 1220, 1070
    draw.text((left, 120), "A　安全—可用性平面", font=panel_font, fill=black)
    draw.rectangle((left, top, right, bottom), outline=black, width=3)

    def px(x: float) -> float:
        return left + x * (right - left)

    def py(y: float) -> float:
        return bottom - y * (bottom - top)

    def point_xy(x: float, y: float) -> tuple[int, int]:
        return (
            int(min(max(px(x), left + 34), right - 34)),
            int(min(max(py(y), top + 34), bottom - 34)),
        )

    draw.rectangle((px(0.8), py(1.0), px(1.0), py(0.8)), fill=green)
    draw.text((px(0.815), py(0.965)), "理想区域", font=small_font, fill="#39724b")
    for value in (0.0, 0.2, 0.4, 0.6, 0.8, 1.0):
        x = px(value)
        y = py(value)
        draw.line((x, top, x, bottom), fill=light_grey, width=2)
        draw.line((left, y, right, y), fill=light_grey, width=2)
        draw_centered(draw, (int(x), bottom + 35), f"{value:.1f}", small_font, grey)
        box = draw.textbbox((0, 0), f"{value:.1f}", font=small_font)
        draw.text((left - 22 - (box[2] - box[0]), y - (box[3] - box[1]) / 2), f"{value:.1f}", font=small_font, fill=grey)

    dashed_line(draw, (px(0), py(1)), (px(1), py(0)), fill="#9ca4aa", width=3, dash=18)
    draw_centered(draw, ((left + right) // 2, bottom + 82), "benign completion（合法完成率）", label_font, black)
    y_label = Image.new("RGBA", (520, 60), (255, 255, 255, 0))
    y_draw = ImageDraw.Draw(y_label)
    draw_centered(y_draw, (260, 30), "attack capture（攻击捕获率）", label_font, black)
    y_label = y_label.rotate(90, expand=True)
    image.paste(y_label, (28, top + 190), y_label)

    # Baseline points.
    for x, y in ((1.0, 0.0), (0.2, 0.8), (0.1, 0.9), (0.0, 1.0)):
        cx, cy = point_xy(x, y)
        draw.ellipse((cx - 17, cy - 17, cx + 17, cy + 17), fill="#9ca4aa", outline=black, width=4)
    # Ablations.
    triangle(draw, point_xy(0.1, 0.9), 24, "#72a9d6", blue)
    triangle(draw, point_xy(0.0, 1.0), 24, "#72a9d6", blue)
    # Full PUB.
    star(draw, point_xy(0.0, 1.0), 34, 15, "#e2676b", red)

    annotations = (
        (point_xy(1.0, 0.0), (right - 300, bottom - 135), "Naked / Prompt\n(1.0, 0.0)"),
        (point_xy(0.2, 0.8), (px(0.31), py(0.72)), "Content SHA\n(0.2, 0.8)"),
        (point_xy(0.1, 0.9), (px(0.25), py(0.95)), "Endpoint / Access\n(0.1, 0.9)"),
        (point_xy(0.0, 1.0), (px(0.28), py(0.86)), "Keyword / Process / PUB\n(0.0, 1.0)"),
    )
    for source, target, text in annotations:
        draw.line((source[0], source[1], target[0], target[1]), fill=grey, width=2)
        draw.multiline_text(target, text, font=tiny_font, fill=black, spacing=3)

    legend_y = 1250
    draw.ellipse((170, legend_y - 12, 194, legend_y + 12), fill="#9ca4aa", outline=black, width=3)
    draw.text((207, legend_y - 17), "baseline", font=small_font, fill=black)
    triangle(draw, (395, legend_y), 16, "#72a9d6", blue, width=3)
    draw.text((425, legend_y - 17), "ablation", font=small_font, fill=black)
    star(draw, (700, legend_y), 21, 9, "#e2676b", red, width=3)
    draw.text((735, legend_y - 17), "Full PUB", font=small_font, fill=black)

    # Right-side F1 panel.
    panel_left, panel_right = 1370, 2110
    draw.text((panel_left, 120), "B　二分类 F1", font=panel_font, fill=black)
    methods = (
        ("Naked", 0.000, grey),
        ("Prompt", 0.000, grey),
        ("Keyword", 0.667, grey),
        ("Content SHA", 0.615, grey),
        ("Endpoint", 0.643, grey),
        ("Access-only", 0.643, blue),
        ("Process-only", 0.667, blue),
        ("Full PUB", 0.667, red),
    )
    label_right = panel_left + 205
    bar_left = panel_left + 230
    bar_right = panel_right - 25
    scale_max = 0.7
    for tick in (0.0, 0.2, 0.4, 0.6):
        x = bar_left + tick / scale_max * (bar_right - bar_left)
        draw.line((x, 190, x, 1065), fill=light_grey, width=2)
        draw_centered(draw, (int(x), 1095), f"{tick:.1f}", small_font, grey)
    for index, (name, value, color) in enumerate(methods):
        y = 245 + index * 102
        name_box = draw.textbbox((0, 0), name, font=small_font)
        draw.text((label_right - (name_box[2] - name_box[0]), y - 14), name, font=small_font, fill=black)
        length = value / scale_max * (bar_right - bar_left)
        if value:
            draw.rounded_rectangle((bar_left, y - 18, bar_left + length, y + 25), radius=9, fill=color)
        draw.text((bar_left + max(length, 8) + 12, y - 14), f"{value:.3f}", font=small_font, fill=black)

    draw.text((panel_left, 1225), "n = 10 attack + 10 control；同一 v3 oracle", font=small_font, fill=grey)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", dpi=(300, 300), optimize=True)


def clone_paragraph_properties(source, target) -> None:
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def clone_run_properties(source_run, target_run) -> None:
    if target_run._r.rPr is not None:
        target_run._r.remove(target_run._r.rPr)
    if source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def add_cloned_text(paragraph, template, text: str):
    clone_paragraph_properties(template, paragraph)
    run = paragraph.add_run(text)
    if template.runs:
        clone_run_properties(template.runs[0], run)
    return run


def replace_text_preserving_first_run(paragraph, text: str) -> None:
    rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)


def append_preserving_last_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    if len(paragraph.runs) > 1:
        clone_run_properties(paragraph.runs[-2], run)


def superscript_citations(document: Document) -> int:
    citation_pattern = re.compile(r"\[\d+\]")
    reference_start = next(
        (index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().startswith("引用源")),
        len(document.paragraphs),
    )
    converted = 0
    for paragraph in document.paragraphs[:reference_start]:
        for run in list(paragraph.runs):
            matches = list(citation_pattern.finditer(run.text))
            if not matches:
                continue
            parent = run._r.getparent()
            position = parent.index(run._r)
            rpr = deepcopy(run._r.rPr) if run._r.rPr is not None else None
            cursor = 0
            pieces: list[tuple[str, bool]] = []
            for match in matches:
                if match.start() > cursor:
                    pieces.append((run.text[cursor:match.start()], False))
                pieces.append((match.group(0), True))
                cursor = match.end()
            if cursor < len(run.text):
                pieces.append((run.text[cursor:], False))
            for text, is_citation in pieces:
                new_run = OxmlElement("w:r")
                if rpr is not None:
                    new_run.append(deepcopy(rpr))
                if is_citation:
                    properties = new_run.get_or_add_rPr()
                    for name in ("w:vertAlign", "w:sz", "w:szCs"):
                        for old in properties.findall(qn(name)):
                            properties.remove(old)
                    vert = OxmlElement("w:vertAlign")
                    vert.set(qn("w:val"), "superscript")
                    properties.append(vert)
                    for name in ("w:sz", "w:szCs"):
                        size = OxmlElement(name)
                        size.set(qn("w:val"), "13")
                        properties.append(size)
                    converted += 1
                text_node = OxmlElement("w:t")
                if text[:1].isspace() or text[-1:].isspace():
                    text_node.set(qn("xml:space"), "preserve")
                text_node.text = text
                new_run.append(text_node)
                parent.insert(position, new_run)
                position += 1
            parent.remove(run._r)
    return converted


def edit_document() -> None:
    make_chart(CHART)
    document = Document(SOURCE)
    paragraphs = document.paragraphs
    heading_template = next(p for p in paragraphs if p.text.strip().startswith("十、验证"))
    body_template = paragraphs[paragraphs.index(heading_template) + 1]
    caption_template = next(p for p in paragraphs if p.text.strip().startswith("【图10插入处"))
    conclusion = next(p for p in paragraphs if p.text.strip().startswith("十一、结论"))

    # Keep the historical smoke statement accurate after the new paired run.
    smoke = next(p for p in paragraphs if "baseline_ablation_recorded=false" in p.text)
    old = "manifest 同时写明 baseline_ablation_recorded=false，因此这轮也没有完成 endpoint-only、去 A、去 S、去 T 的不同情况。"
    new = "原 smoke manifest 写明 baseline_ablation_recorded=false；补跑的同题 baseline 与消融结果列于第十一节。"
    if old not in smoke.text:
        raise RuntimeError("expected historical baseline sentence not found")
    replace_text_preserving_first_run(smoke, smoke.text.replace(old, new))

    evidence_intro = next(p for p in paragraphs if "主 SHA256SUMS 共列出 23 个文件" in p.text)
    append_preserving_last_run(
        evidence_intro,
        " 新增 baseline run 由 BASELINE_RUN_SHA256SUMS 与 BASELINE_SCORE_SHA256SUMS 单独封存。",
    )

    heading = conclusion.insert_paragraph_before()
    add_cloned_text(heading, heading_template, TITLE)
    paragraph_1 = conclusion.insert_paragraph_before()
    add_cloned_text(paragraph_1, body_template, BODY_1)

    figure_paragraph = conclusion.insert_paragraph_before()
    clone_paragraph_properties(body_template, figure_paragraph)
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.first_line_indent = Pt(0)
    figure_paragraph.paragraph_format.left_indent = Pt(0)
    figure_paragraph.paragraph_format.right_indent = Pt(0)
    figure_paragraph.paragraph_format.line_spacing = 1
    figure_paragraph.paragraph_format.space_before = Pt(6)
    figure_paragraph.paragraph_format.space_after = Pt(4)
    figure_paragraph.paragraph_format.keep_with_next = True
    figure_paragraph.paragraph_format.keep_together = True
    picture_run = figure_paragraph.add_run()
    shape = picture_run.add_picture(str(CHART), width=Cm(14.35))
    shape._inline.docPr.set(
        "descr",
        "Baseline and PUB security-utility plane with attack capture, benign completion, and binary F1.",
    )
    shape._inline.docPr.set("title", "Baseline 与 PUB 安全—可用性对照")

    caption = conclusion.insert_paragraph_before()
    add_cloned_text(caption, caption_template, CAPTION)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_together = True

    paragraph_2 = conclusion.insert_paragraph_before()
    add_cloned_text(paragraph_2, body_template, BODY_2)
    paragraph_3 = conclusion.insert_paragraph_before()
    add_cloned_text(paragraph_3, body_template, BODY_3)

    replace_text_preserving_first_run(conclusion, conclusion.text.replace("十一、结论", "十二、结论", 1))
    converted = superscript_citations(document)
    if converted != 8:
        raise RuntimeError(f"expected 8 body citations, converted {converted}")

    document.save(OUTPUT)
    print(f"output={OUTPUT}")
    print(f"chart={CHART}")
    print(f"baseline_body_chars={len(BODY_1) + len(BODY_2) + len(BODY_3)}")
    print(f"citations_superscripted={converted}")


if __name__ == "__main__":
    edit_document()
